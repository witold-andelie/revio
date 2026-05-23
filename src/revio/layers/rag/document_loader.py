"""Document loading for RAG indexing.

Supports the typical "company guideline" formats:
- Markdown (.md) — split by headers
- Plain text (.txt, .rst, .adoc) — chunked by paragraph
- PDF (.pdf) — extracted page by page, chunked
- DOCX (.docx) — split by heading paragraphs

We deliberately do NOT load source code files into the RAG index — the agent
already reads code directly via `read_file`/`get_function_at`. RAG is for
the *meta-text* describing how code should be reviewed.

Adapted from v1's src/rag/document_loader.py, simplified for guideline-only use.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from langchain_core.documents import Document


logger = logging.getLogger(__name__)


# Supported guideline extensions. We don't try to load code files into RAG.
DOC_EXTENSIONS = frozenset({".md", ".markdown", ".txt", ".rst", ".adoc", ".pdf", ".docx"})


class DocumentLoader:
    """Load and chunk guideline documents for vector indexing."""

    DEFAULT_CHUNK_CHARS = 1000

    # --- Public entry points --------------------------------------------------

    @classmethod
    def load_file(cls, file_path: str | Path) -> list[Document]:
        """Load a single file. Dispatches by extension."""
        path = Path(file_path).expanduser().resolve()
        if not path.is_file():
            logger.warning("DocumentLoader: not a file: %s", path)
            return []

        ext = path.suffix.lower()
        if ext in (".md", ".markdown"):
            return cls._load_markdown(path)
        if ext in (".txt", ".rst", ".adoc"):
            return cls._load_text(path)
        if ext == ".pdf":
            return cls._load_pdf(path)
        if ext == ".docx":
            return cls._load_docx(path)

        logger.warning("DocumentLoader: unsupported extension %s for %s", ext, path)
        return []

    @classmethod
    def load_directory(
        cls,
        dir_path: str | Path,
        *,
        max_files: int = 500,
        ignore_hidden: bool = True,
    ) -> list[Document]:
        """Recursively load all guideline files from a directory."""
        root = Path(dir_path).expanduser().resolve()
        if not root.is_dir():
            return []

        documents: list[Document] = []
        count = 0
        for path in sorted(root.rglob("*")):
            if count >= max_files:
                logger.warning("DocumentLoader: reached max_files (%d), stopping", max_files)
                break
            if not path.is_file():
                continue
            if ignore_hidden and any(part.startswith(".") for part in path.relative_to(root).parts):
                continue
            if path.suffix.lower() not in DOC_EXTENSIONS:
                continue
            docs = cls.load_file(path)
            if docs:
                documents.extend(docs)
                count += 1

        return documents

    # --- Format-specific loaders ---------------------------------------------

    @classmethod
    def _load_markdown(cls, path: Path) -> list[Document]:
        try:
            content = path.read_text(encoding="utf-8", errors="ignore")
        except OSError as e:
            logger.warning("read failed %s: %s", path, e)
            return []

        sections = cls._split_by_headers(content)
        docs: list[Document] = []
        for i, (title, body) in enumerate(sections):
            body = body.strip()
            if not body:
                continue
            # Chunk long sections; keep section title in metadata
            for j, chunk in enumerate(cls._chunk_text(body, cls.DEFAULT_CHUNK_CHARS)):
                docs.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": str(path),
                        "section_title": title,
                        "section_index": i,
                        "chunk_index": j,
                        "doc_type": "guideline",
                        "format": "markdown",
                    },
                ))
        return docs

    @classmethod
    def _load_text(cls, path: Path) -> list[Document]:
        try:
            content = path.read_text(encoding="utf-8", errors="ignore")
        except OSError as e:
            logger.warning("read failed %s: %s", path, e)
            return []

        docs: list[Document] = []
        for i, chunk in enumerate(cls._chunk_text(content, cls.DEFAULT_CHUNK_CHARS)):
            docs.append(Document(
                page_content=chunk,
                metadata={
                    "source": str(path),
                    "chunk_index": i,
                    "doc_type": "guideline",
                    "format": path.suffix.lstrip(".") or "text",
                },
            ))
        return docs

    @classmethod
    def _load_pdf(cls, path: Path) -> list[Document]:
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("pypdf not installed; cannot load PDF: %s", path)
            return []

        try:
            reader = PdfReader(str(path))
        except Exception as e:
            logger.warning("PDF parse failed %s: %s", path, e)
            return []

        docs: list[Document] = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if not text.strip():
                continue
            for j, chunk in enumerate(cls._chunk_text(text, cls.DEFAULT_CHUNK_CHARS)):
                docs.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": str(path),
                        "page_number": page_num + 1,
                        "chunk_index": j,
                        "doc_type": "guideline",
                        "format": "pdf",
                    },
                ))
        return docs

    @classmethod
    def _load_docx(cls, path: Path) -> list[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("python-docx not installed; cannot load DOCX: %s", path)
            return []

        try:
            d = DocxDocument(str(path))
        except Exception as e:
            logger.warning("DOCX parse failed %s: %s", path, e)
            return []

        docs: list[Document] = []
        current_section = "Document"
        buffer: list[str] = []

        def flush(section: str):
            if not buffer:
                return
            text = "\n".join(buffer)
            buffer.clear()
            for j, chunk in enumerate(cls._chunk_text(text, cls.DEFAULT_CHUNK_CHARS)):
                docs.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": str(path),
                        "section_title": section,
                        "chunk_index": j,
                        "doc_type": "guideline",
                        "format": "docx",
                    },
                ))

        for para in d.paragraphs:
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading"):
                flush(current_section)
                current_section = para.text.strip() or current_section
            elif para.text.strip():
                buffer.append(para.text)

        flush(current_section)
        return docs

    # --- Helpers --------------------------------------------------------------

    @staticmethod
    def _chunk_text(text: str, max_chars: int) -> list[str]:
        """Greedy paragraph-aware chunker."""
        paragraphs = text.split("\n\n")
        chunks: list[str] = []
        current = ""
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) + 2 > max_chars and current:
                chunks.append(current)
                current = para
            else:
                current = current + "\n\n" + para if current else para
        if current:
            chunks.append(current)
        if not chunks:
            chunks = [text[:max_chars]] if text else []
        return chunks

    @staticmethod
    def _split_by_headers(content: str) -> list[tuple[str, str]]:
        """Split markdown by ATX headers, preserving structure."""
        pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
        matches = list(pattern.finditer(content))

        if not matches:
            return [("Document", content)]

        sections: list[tuple[str, str]] = []
        # Preamble before first header
        if matches[0].start() > 0:
            preamble = content[:matches[0].start()].strip()
            if preamble:
                sections.append(("Introduction", preamble))

        for i, match in enumerate(matches):
            title = match.group(2).strip()
            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            body = content[start:end].strip()
            sections.append((title, body))

        return sections
